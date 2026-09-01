"""Exporta texto em markdown (pareceres, minutas) para Word (.docx).

O usuário recebia um arquivo .md que o Windows não sabe abrir — "tive que
escolher o programa". Word abre em qualquer máquina de licitação e de lá
vira PDF com um clique. A conversão cobre o markdown que os nossos
prompts realmente produzem: títulos, listas, citações, tabelas simples e
negrito/itálico inline. O que não for reconhecido sai como parágrafo.
"""
import io
import re

_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def _escrever_inline(paragrafo, texto):
    """Aplica **negrito**, *itálico* e `código` como runs do parágrafo."""
    for parte in _INLINE.split(texto):
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**"):
            paragrafo.add_run(parte[2:-2]).bold = True
        elif parte.startswith("*") and parte.endswith("*") and len(parte) > 2:
            paragrafo.add_run(parte[1:-1]).italic = True
        elif parte.startswith("`") and parte.endswith("`"):
            run = paragrafo.add_run(parte[1:-1])
            run.font.name = "Consolas"
        else:
            paragrafo.add_run(parte)


def _e_linha_de_tabela(linha):
    return linha.startswith("|") and linha.rstrip().endswith("|")


def _celulas(linha):
    return [c.strip() for c in linha.strip().strip("|").split("|")]


def markdown_para_docx(texto, titulo=None):
    """Converte o markdown num .docx e devolve os bytes prontos."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)
    if titulo:
        doc.add_heading(titulo, level=0)

    linhas = (texto or "").splitlines()
    i = 0
    while i < len(linhas):
        linha = linhas[i].rstrip()
        crua = linha.strip()
        if not crua:
            i += 1
            continue
        if crua.startswith("#"):
            nivel = min(len(crua) - len(crua.lstrip("#")), 4)
            doc.add_heading(crua.lstrip("#").strip(), level=nivel)
        elif crua.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            _escrever_inline(p, crua.lstrip("> ").strip())
        elif re.match(r"^[-*+]\s+", crua):
            p = doc.add_paragraph(style="List Bullet")
            _escrever_inline(p, re.sub(r"^[-*+]\s+", "", crua))
        elif re.match(r"^\d+[.)]\s+", crua):
            p = doc.add_paragraph(style="List Number")
            _escrever_inline(p, re.sub(r"^\d+[.)]\s+", "", crua))
        elif crua in ("---", "***", "___"):
            doc.add_paragraph()
        elif _e_linha_de_tabela(crua):
            bloco = []
            while i < len(linhas) and _e_linha_de_tabela(linhas[i].strip()):
                celulas = _celulas(linhas[i])
                # linha separadora |---|---| não vira linha da tabela
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in celulas):
                    bloco.append(celulas)
                i += 1
            if bloco:
                colunas = max(len(l) for l in bloco)
                tabela = doc.add_table(rows=len(bloco), cols=colunas)
                tabela.style = "Light Grid Accent 1"
                for r, valores in enumerate(bloco):
                    for c in range(colunas):
                        celula = tabela.cell(r, c).paragraphs[0]
                        _escrever_inline(celula,
                                         valores[c] if c < len(valores)
                                         else "")
            continue
        else:
            p = doc.add_paragraph()
            _escrever_inline(p, crua)
        i += 1

    saida = io.BytesIO()
    doc.save(saida)
    return saida.getvalue()


MEDIA_DOCX = ("application/vnd.openxmlformats-officedocument"
              ".wordprocessingml.document")
